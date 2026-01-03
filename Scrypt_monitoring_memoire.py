import pandas as pd
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import TimeoutException, NoSuchElementException, WebDriverException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
import time
from docx import Document
import win32com.client
import os
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from datetime import datetime
import pyodbc
from dotenv import load_dotenv

# --- OPTIONS CHROME ---
chrome_options = Options()
chrome_options.add_argument("--headless=new")  # headless moderne (Chrome 109+)
# Les deux options ci-dessous sont surtout utiles en Linux/CI; inoffensives sous Windows :
chrome_options.add_argument("--no-sandbox")
chrome_options.add_argument("--disable-dev-shm-usage")

# --- LANCEMENT SANS CHEMIN DE DRIVER ---
# >>> Laisse Selenium Manager télécharger la bonne version du driver automatiquement <<<
driver = webdriver.Chrome(options=chrome_options)

def click_cookies_if_any(driver):
    selectors = [
        "#onetrust-accept-btn-handler",   # très courant
        ".ot-sdk-container .accept-btn",
        "button[aria-label='Accepter']",
        "button.cookie-accept",
    ]
    for css in selectors:
        try:
            btn = WebDriverWait(driver, 2).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, css))
            )
            btn.click()
            break  # si trouvé, on s'arrête
        except Exception:
            pass

def scrape_pharma_gdd(url):
    # Vérifie si l'URL est "Pas d'URL"
    if url == "Pas d'URL":
        return "Pas d'URL"
    
    try:
        driver.get(url)
        time.sleep(3)  
        soup = BeautifulSoup(driver.page_source, 'html.parser')  
        
        price_element = soup.select_one("#product-adding-to-cart > div.variantOptionsAndCart > div.d-flex.addtocart > div.price-infos > div > span.price")
        
        if price_element:
            return price_element.text.strip()
        else:
            return "Prix non trouvé"
    except Exception as e:
        print(f"Erreur pour {url}: {e}")
        return "Erreur de récupération"
    
def scrape_pharmashopi(url):
    if not isinstance(url, str) or not url.startswith("http"):
        return "Pas d'URL"

    try:
        driver.get(url)
        time.sleep(3)

        soup = BeautifulSoup(driver.page_source, 'html.parser')
        price_element = soup.select_one("span.price")  # ✅ Sélecteur stable

        if price_element:
            return price_element.text.strip().replace('\xa0€', ' €')
        else:
            return "Prix non trouvé"

    except Exception as e:
        print(f"Erreur pour {url}: {e}")
        return "Erreur de récupération"
    
def scrape_redcare_pharmacie(url):
    if url == "Pas d'URL":
        return "Pas d'URL"
    
    try:
        driver.get(url)
        click_cookies_if_any(driver)   # <--- ajout ici

        wait = WebDriverWait(driver, 10)
        price_element = wait.until(EC.presence_of_element_located((
            By.XPATH, "/html/body/main/div[6]/div[1]/div[3]/div[2]/div[1]/div[2]/div[1]/div[2]/div[1]/div[1]"
        )))
        
        if price_element:
            return price_element.text.strip()
        else:
            return "Prix non trouvé"
    except TimeoutException:
        print(f"Timeout : le prix n'a pas chargé à temps pour {url}")
        return "Timeout"
    except NoSuchElementException:
        print(f"Erreur : élément prix non trouvé sur {url}")
        return "Prix non trouvé"
    except Exception as e:
        print(f"Erreur inattendue pour {url}: {e}")
        return "Erreur de récupération"
def scrape_universpharmacie(url):
    # Vérifie si l'URL est "Pas d'URL"
    if url == "Pas d'URL":
        return "Pas d'URL"
    
    try:
        driver.get(url)
        time.sleep(3)  # Attendre que la page se charge complètement
        soup = BeautifulSoup(driver.page_source, 'html.parser')  # Créer l'objet BeautifulSoup
        
        # Utilisation du sélecteur CSS pour trouver le prix
        price_element = soup.select_one("#main > div.row.product-container > div:nth-child(2) > div.product-prices > div.product-price.h5 > div > span")
        
        if price_element:
            return price_element.text.strip()
        else:
            return "Prix non trouvé"
    except Exception as e:
        print(f"Erreur pour {url}: {e}")
        return "Erreur de récupération"
def scrape_boticinal(url):
    # Vérifie si l'URL est "Pas d'URL"
    if url == "Pas d'URL":
        return "Pas d'URL"
    
    try:
        driver.get(url)
        time.sleep(3)  # Attendre que la page se charge complètement
        soup = BeautifulSoup(driver.page_source, 'html.parser')  # Créer l'objet BeautifulSoup
        
        # Tentatives multiples pour récupérer le prix
        price_element = soup.select_one(".price, .product-price, .product-price-regular")
        if not price_element:
            price_element = soup.find("span", {"data-price": True})
        
        if price_element:
            return price_element.text.strip()
        else:
            return "Prix non trouvé"
    except Exception as e:
        print(f"Erreur pour {url}: {e}")
        return "Erreur de récupération"
def scrape_pharmacie_citypharma(url):
    # Vérifie si l'URL est "Pas d'URL"
    if url == "Pas d'URL":
        return "Pas d'URL"
    
    try:
        driver.get(url)
        time.sleep(3)  
        soup = BeautifulSoup(driver.page_source, 'html.parser')  
        
        price_element = soup.select_one("#product_header > div.col-md-6.content_product > div > div.product-information > div.product-prices > div.product-price.h5 > div > span:nth-child(2)")
        
        if price_element:
            return price_element.text.strip()
        else:
            return "Prix non trouvé"
    except Exception as e:
        print(f"Erreur pour {url}: {e}")
        return "Erreur de récupération"
def pharmacie_polygone(url):
    # Vérifie si l'URL est "Pas d'URL"
    if url == "Pas d'URL":
        return "Pas d'URL"
    
    try:
        driver.get(url)
        time.sleep(3)  
        soup = BeautifulSoup(driver.page_source, 'html.parser')  
        
        price_element = soup.select_one("#appContainer > div > div.product-show > div > div:nth-child(3) > div.col-md-9.col-lg-7.product-content > div:nth-child(6) > div > div.price-container > div > span > span")
        
        if price_element:
            return price_element.text.strip()
        else:
            return "Prix non trouvé"
    except Exception as e:
        print(f"Erreur pour {url}: {e}")
        return "Erreur de récupération"
def scrape_pharmashopdiscount(url):
    # Vérifie si l'URL est "Pas d'URL"
    if url == "Pas d'URL":
        return "Pas d'URL"
    
    try:
        driver.get(url)
        click_cookies_if_any(driver)
        time.sleep(3)  # Attendre que la page se charge complètement
        soup = BeautifulSoup(driver.page_source, 'html.parser')  # Créer l'objet BeautifulSoup
        
        # Récupération du prix en promotion (le prix actuel)
        promo_price_element = soup.select_one("#price")
        
        # Récupération du prix barré (le prix de base)
        base_price_element = soup.select_one("body > main > div:nth-child(2) > div.col-lg-8.mt-3 > div.product-cart-add > p > s")
        
        # Retourner le prix promotionnel s'il existe, sinon le prix barré
        if promo_price_element:
            return promo_price_element.text.strip()
        elif base_price_element:
            return base_price_element.text.strip()
        else:
            return "Prix non trouvé"
    except Exception as e:
        print(f"Erreur pour {url}: {e}")
        return "Erreur de récupération"
def scrape_parapharmalafayette(url):
    # Vérifie si l'URL est "Pas d'URL"
    if url == "Pas d'URL":
        return "Pas d'URL"
    
    try:
        driver.get(url)
        click_cookies_if_any(driver)
        time.sleep(3)  # Attendre que la page se charge complètement
        
        # Attente explicite pour que le prix soit visible
        new_price_element = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, "/html/body/div[2]/main/div[1]/div/div/div/div[1]/div/div[2]/div[1]/div[3]/div[2]/div[1]/span/span/span[1]"))
        )

        if new_price_element:
            # Si un nouveau prix est trouvé, on le retourne
            return new_price_element.text.strip()
        else:
            return "Prix non trouvé"
    except Exception as e:
        print(f"Erreur pour {url}: {e}")
        return "Erreur de récupération"
    
def scrape_parapharmadirect(url):
    if not isinstance(url, str) or not url.startswith("http"):
        return "Pas d'URL"
    try:
        driver.get(url)
        time.sleep(3)  # Attendre que la page se charge complètement
        soup = BeautifulSoup(driver.page_source, 'html.parser')
        # Nouveau sélecteur plus générique et stable
        price_element = soup.select_one("div.product-price span")
        if price_element:
            return price_element.text.strip()
        else:
            return "Prix non trouvé"
    except Exception as e:
        print(f"Erreur pour {url}: {e}")
        return "Erreur de récupération"

    
def scrape_pharmacie3000(url):
    if not isinstance(url, str) or not url.startswith("http"):
        return "Pas d'URL"
    
    try:
        driver.get(url)
        time.sleep(3)  # Attendre que la page se charge complètement
        
        # Utilisation de Selenium avec XPath direct
        price_element = driver.find_element(By.XPATH, '//*[@id="our_price_display"]')
        
        return price_element.text.strip() if price_element else "Prix non trouvé"
    
    except Exception as e:
        print(f"Erreur pour {url}: {e}")
        return "Erreur de récupération"

def scrape_pharma360(url):
    # Vérifie si l'URL est "Pas d'URL"
    if url == "Pas d'URL":
        return "Pas d'URL"
    
    try:
        driver.get(url)
        time.sleep(3)  
        soup = BeautifulSoup(driver.page_source, 'html.parser')  

        # Sélecteurs pour le prix réduit et le prix de base
        reduced_price_selector = "#add-to-cart-or-refresh > div.page-product__right__grid > div > div > div > div.product-add-to-cart__prices > div > div.product-price.has-discount > div > span:nth-child(1)"
        base_price_selector = "#add-to-cart-or-refresh > div.page-product__right__grid > div > div > div > div.product-add-to-cart__prices > div > div > div > span"
        
        # Trouver le prix réduit
        reduced_price_element = soup.select_one(reduced_price_selector)
        
        # Si le prix réduit existe, le retourner
        if reduced_price_element:
            return reduced_price_element.text.strip()
        
        # Sinon, chercher le prix de base
        base_price_element = soup.select_one(base_price_selector)
        if base_price_element:
            return base_price_element.text.strip()
        
        return "Prix non trouvé"
    except Exception as e:
        print(f"Erreur pour {url}: {e}")
        return "Erreur de récupération"
def scrape_pharmaciedesdrakkars(url):
    # Vérifie si l'URL est "Pas d'URL"
    if url == "Pas d'URL":
        return "Pas d'URL"
    
    try:
        driver.get(url)
        time.sleep(3)  
        soup = BeautifulSoup(driver.page_source, 'html.parser')  

        price_selector = "#product_price > span > strong"
        price_element = soup.select_one(price_selector)
        
        # Si le prix est trouvé, le retourner
        if price_element:
            return price_element.text.strip()
        
        return "Prix non trouvé"
    except Exception as e:
        print(f"Erreur pour {url}: {e}")
        return "Erreur de récupération"
def scrape_parapharmacie_et_medicament(url):
    # Vérifie si l'URL est "Pas d'URL"
    if url == "Pas d'URL":
        return "Pas d'URL"
    
    try:
        driver.get(url)
        time.sleep(3)  
        soup = BeautifulSoup(driver.page_source, 'html.parser')  

        price_selector = "#form_add_to_cart_product_page > fieldset > div.c_add_to_cart_area > div.row.row-cols-auto.justify-content-between.align-items-center.mb-3 > div:nth-child(1) > div.h1.text-primary.mb-0"
        price_element = soup.select_one(price_selector)

        if price_element:
            return price_element.text.strip().replace('\xa0€', ' €')
        
        return "Prix non trouvé"
    except Exception as e:
        print(f"Erreur pour {url}: {e}")
        return "Erreur de récupération"


# Dictionnaire associant chaque site à sa fonction
site_functions = {
    "pharma-gdd": scrape_pharma_gdd,
    "pharmashopi": scrape_pharmashopi,
    "redcare-pharmacie": scrape_redcare_pharmacie,
    "universpharmacie": scrape_universpharmacie,
    "boticinal": scrape_boticinal,
    "pharmacie-citypharma": scrape_pharmacie_citypharma,
    "pharmaciepolygone": pharmacie_polygone,
    "pharmashopdiscount": scrape_pharmashopdiscount,
    "parapharmalafayette": scrape_parapharmalafayette,
    "parapharmadirect": scrape_parapharmadirect,
    "pharmacie3000": scrape_pharmacie3000,
    "pharma360": scrape_pharma360, 
    "pharmaciedesdrakkars": scrape_pharmaciedesdrakkars,
    "parapharmacie et medicament": scrape_parapharmacie_et_medicament,
}

# Définir le dossier de sortie pour éviter les problèmes de chemin
output_dir = R"C:\Users\HugoMAIROT\OneDrive - GTH\Bureau\Sujet_Fevrier25\Scraping_verif_prix-remi"

# Obtenir la date actuelle au format JJMMAA (ex: "150225" pour le 15 février 2025)
date_format = datetime.now().strftime("%d%m%y")

file_path = R"C:\Users\HugoMAIROT\OneDrive - GTH\Bureau\Sujet_Fevrier25\Scraping_verif_prix-remi\Produits cibles pharmacies Hybrides.xlsx"  # Remplacez par le chemin vers votre fichier
output_path = os.path.join(output_dir, f"Alertes_Prix_{date_format}.xlsx")

# Historique CSV
history_csv_path = os.path.join(output_dir, "Historique_Prix.csv")

# SQL
load_dotenv()
email = os.getenv("SQL_EMAIL")
password = os.getenv("SQL_PASSWORD")

TABLE_NAME = "dbo.PriceMonitoringHistorique"  # <-- adapte au nom exact de ta table


df = pd.read_excel(file_path)

## Parcours des colonnes de sites
for site, func in site_functions.items():
    if site in df.columns:
        print(f"Scraping pour le site : {site}")
        df[site] = df[site].apply(
            lambda url: func(url) if isinstance(url, str) and url.startswith("http") else "Pas d'URL"
        )

# Identifier la colonne du seuil "PPC TTC min"
col_seuil = "PPC TTC min"

# Vérifier que la colonne existe
if col_seuil in df.columns:
    # Fonction de nettoyage des prix
    def nettoyer_prix(val):
        if isinstance(val, str):
            val = val.strip()  # Supprimer les espaces
            
            # Vérifier si la valeur est une exception à garder
            if val.lower() in ["prix non trouvé", "Pas d'URL", ""]:
                return val  # On la garde telle quelle
            
            val = val.replace(",", ".")  # Remplacer les virgules par des points
            val = ''.join(c for c in val if c.isdigit() or c == ".")  # Garder uniquement les chiffres et points
            
            try:
                return float(val)  # Convertir en float
            except ValueError:
                return None  # Si la conversion échoue, mettre NaN

    # Liste des colonnes à exclure du nettoyage
    colonnes_a_exclure = ["DÉSIGNATION", "PPC TTC min", "Nombre de site ayant le produit", "PPC TTC", "EAN", "easypara"]

    # Sauvegarder l'ordre des colonnes d'origine
    ordre_colonnes = df.columns

    # Nettoyer uniquement les colonnes concernées
    df_clean = df.drop(columns=colonnes_a_exclure).applymap(nettoyer_prix)

    # Remettre les colonnes nettoyées et exclues en respectant l'ordre d'origine
    df.update(df_clean)  # Met à jour uniquement les colonnes nettoyées, sans toucher aux autres
    df = df[ordre_colonnes]  # Réorganiser les colonnes dans leur ordre initial

    # Définir les colonnes à exclure des calculs
    colonnes_a_exclure_calculs = ["PPC TTC", "EAN","easypara","Nombre de site ayant le produit", col_seuil]

    # Transformer en float uniquement les colonnes contenant des prix
    df_numeric = df.apply(pd.to_numeric, errors="coerce")

    # Créer un dictionnaire pour stocker le nombre de produits sous le seuil par site
    seuil_counts = {}

    # Parcourir chaque colonne sauf celles à exclure
    for col in df_numeric.columns:
        if col not in colonnes_a_exclure_calculs:
            seuil_counts[col] = (df_numeric[col] < df_numeric[col_seuil]).sum()

    # Ajouter la ligne des totaux au DataFrame
    df.loc["Total sous seuil"] = seuil_counts
    df.loc["Total sous seuil", "DÉSIGNATION"] = "Nombre de produits sous le seuil"  # Nom explicite

    # Créer un dictionnaire pour stocker le nombre total de produits trouvés par site (hors "Non trouvé")
    produits_trouves_counts = {
        col: df_numeric[col].notna().sum() - (df[col].astype(str) == "Non trouvé").sum()
        for col in df_numeric.columns if col not in colonnes_a_exclure_calculs
    }

    # Créer un dictionnaire pour stocker le pourcentage de produits sous le seuil par site
    pourcentage_sous_seuil = {}

    for col in seuil_counts:
        total_produits_trouves = produits_trouves_counts.get(col, 1)  # Évite une division par zéro
        pourcentage_sous_seuil[col] = f"{round((seuil_counts[col] / total_produits_trouves * 100))}%" if total_produits_trouves > 0 else "0%"

    # Ajouter la ligne du pourcentage sous seuil au DataFrame
    df.loc["% sous seuil"] = pourcentage_sous_seuil
    df.loc["% sous seuil", "DÉSIGNATION"] = "Pourcentage de produits sous le seuil"  # Nom explicite

def add_date_column(df_in: pd.DataFrame, date_col: str = "Date") -> pd.DataFrame:
    df_out = df_in.copy()
    # Date du jour (même format que ton exemple 06/02/2025)
    df_out[date_col] = datetime.now().strftime("%d/%m/%Y")
    return df_out

def remove_summary_rows(df_in: pd.DataFrame) -> pd.DataFrame:
    """Retire les lignes ajoutées pour les stats (Total sous seuil, % sous seuil) si elles existent."""
    df_out = df_in.copy()
    df_out = df_out[df_out["DÉSIGNATION"].notna()]
    df_out = df_out[~df_out["DÉSIGNATION"].isin([
        "Nombre de produits sous le seuil",
        "Pourcentage de produits sous le seuil"
    ])]
    # Ton code ajoute aussi parfois index 'Total sous seuil' / '% sous seuil' : on sécurise
    df_out = df_out[~df_out.index.astype(str).isin(["Total sous seuil", "% sous seuil"])]
    return df_out

def parse_price_to_float(v):
    """Convertit un prix texte en float; renvoie None si non exploitable."""
    if v is None:
        return None
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip()
    if s == "" or s.lower() in ["prix non trouvé", "pas d'url", "erreur de récupération", "timeout"]:
        return None
    s = s.replace("€", "").replace("\xa0", " ").strip()
    s = s.replace(",", ".")
    # garde chiffres/point uniquement
    s = "".join(c for c in s if c.isdigit() or c == ".")
    try:
        return float(s)
    except ValueError:
        return None

def append_to_history_csv(df_week: pd.DataFrame, history_path: str) -> None:
    """Append des lignes à l'historique CSV (créé le fichier si nécessaire)."""
    write_header = not os.path.exists(history_path)
    df_week.to_csv(history_path, mode="a", header=write_header, index=False, sep=";", encoding="utf-8-sig")

def insert_into_sql(df_week: pd.DataFrame, conn_str: str, table_name: str) -> None:
    """
    Insère df_week dans SQL.
    Hypothèse: la table SQL a exactement les mêmes colonnes (mêmes noms) que df_week.
    Si ce n'est pas le cas, il faudra faire un mapping.
    """
    # 1) Ne pas modifier df_week original (sécurité)
    df_sql = df_week.copy()

    # 2) Supprime les colonnes qui ne doivent pas aller en base
    cols_to_drop = ["Nombre de site ayant le produit", "easypara"]
    df_sql = df_sql.drop(columns=[c for c in cols_to_drop if c in df_sql.columns], errors="ignore")

    # 3) Retire les lignes de synthèse si elles traînent (double sécurité)
    if "DÉSIGNATION" in df_sql.columns:
        df_sql = df_sql[df_sql["DÉSIGNATION"].notna()]
        df_sql = df_sql[~df_sql["DÉSIGNATION"].isin([
            "Nombre de produits sous le seuil",
            "Pourcentage de produits sous le seuil"
        ])]

    # Colonnes prix (toutes sauf ces colonnes de référence)
    cols_ref = {"DÉSIGNATION", "PPC TTC min", "PPC TTC", "EAN", "Date"}
    price_cols = [c for c in df_sql.columns if c not in cols_ref]

    # Conversions
    df_sql["PPC TTC min"] = df_sql["PPC TTC min"].apply(parse_price_to_float)
    df_sql["PPC TTC"] = df_sql["PPC TTC"].apply(parse_price_to_float)

    # EAN parfois lu en float/scientifique depuis Excel : on force string propre
    if "EAN" in df_sql.columns:
        df_sql["EAN"] = df_sql["EAN"].apply(lambda x: None if pd.isna(x) else str(int(x)) if isinstance(x, float) else str(x))

    for c in price_cols:
        df_sql[c] = df_sql[c].apply(parse_price_to_float)

    # Remplace NaN pandas par None python
    df_sql = df_sql.where(pd.notnull(df_sql), None)

    # Connexion + insertion bulk
    conn = pyodbc.connect(conn_str)
    try:
        cursor = conn.cursor()
        cursor.fast_executemany = True

        cols = list(df_sql.columns)
        # Attention: noms de colonnes avec espaces/accents -> on encadre avec []
        col_sql = ", ".join([f"[{c}]" for c in cols])
        placeholders = ", ".join(["?"] * len(cols))
        sql = f"INSERT INTO {table_name} ({col_sql}) VALUES ({placeholders})"

        rows = df_sql.values.tolist()
        cursor.executemany(sql, rows)
        conn.commit()
        print(f"✅ Insert SQL OK : {len(rows)} lignes insérées dans {table_name}")

    finally:
        conn.close()



# Sauvegarder les résultats dans le fichier Excel
df.to_excel(output_path, index=False)
print(f"Résultats enregistrés dans : {output_path}")
# --- Création du "batch" hebdo (sans lignes de synthèse) + date ---
df_week = remove_summary_rows(df)
df_week = add_date_column(df_week, "Date")

# --- Append CSV historisé ---
append_to_history_csv(df_week, history_csv_path)
print(f"✅ Historique CSV enrichi : {history_csv_path}")

# --- Connexion SQL ---
conn_str = (
    f"DRIVER={{ODBC Driver 17 for SQL Server}};"
    f"SERVER=sqlserver-dataplatformprd.database.windows.net;"
    f"DATABASE=sqldb-dataplatformprd;"
    f"UID={email};"
    f"PWD={password};"
    "Authentication=ActiveDirectoryPassword"
)

# --- Insert dans SQL ---
insert_into_sql(df_week, conn_str, TABLE_NAME)


##############################################################################################################################################################

# Créer un document Word
doc = Document()
doc.add_heading("Alertes Prix - Produits sous le seuil", level=1)

# Liste pour stocker les alertes
alertes = []

# Vérification des seuils d'alerte
for index, row in df.iterrows():
    produit = row["DÉSIGNATION"]  # Assurez-vous que cette colonne existe
    seuil = row["PPC TTC min"]  # Colonne contenant le seuil d'alerte
    Prix_vente = row["PPC TTC"]
    
    # Liste pour stocker les sites sous le seuil
    sites_sous_seuil = []
    
    # Vérifier chaque site
    for site in df.columns:
        if site not in ["DÉSIGNATION", "PPC TTC min","Nombre de site ayant le produit","PPC TTC","EAN"]:  # Exclure les colonnes non concernées
            prix = row[site]
            
            if prix is not None:  # Vérifie que le prix n'est pas None
                try:
                    prix = float(str(prix).replace("€", "").replace(",", ".").strip())
                    if prix < seuil:
                        sites_sous_seuil.append(f"{site} ({prix}€)")
                except ValueError:
                    continue  # Ignore si la conversion échoue
    
    # Si des sites sont sous le seuil, ajouter une alerte
    if sites_sous_seuil:
        message = f"⚠️ Attention ! Le produit '{produit}' vendu ({Prix_vente}€), est en dessous du seuil sur les sites :\n" + "\n".join(sites_sous_seuil) + "."
        alertes.append(message)
        doc.add_paragraph(message)

# Sauvegarde du fichier Word
output_filename = os.path.join(output_dir, f"Alertes_Prix_{date_format}.docx")
doc.save(output_filename)

# Affichage du statut
if alertes:
    print(f"✅ {len(alertes)} alertes enregistrées dans '{output_filename}'.")
else:
    print("✅ Aucun produit en dessous du seuil d'alerte.")

# # Vérifier si les fichiers existent avant de les attacher
# if not os.path.exists(output_path) or not os.path.exists(output_filename):
#     print("❌ Erreur : Un des fichiers à envoyer n'existe pas.")
#     raise FileNotFoundError("Un des fichiers n'a pas été trouvé.")

try:
    outlook = win32com.client.Dispatch("Outlook.Application")
    mail = outlook.CreateItem(0)  # Crée un nouvel e-mail

    # mail.To = "hmairot@gsa-healthcare.com; rabid@groupe-gth.com; arascle@synergia.eu"
    mail.To = "hmairot@gsa-healthcare.com"

    mail.Subject = "Alertes Prix - Produits sous le seuil"
    mail.Body = "Bonjour à tous,\n\nVous trouverez en pièce jointe le denier rapport des produits dont les prix sont sous le seuil.\nLa dernière ligne indique le total de produits qui se trouvent sous ce seuil sur chaque site.\n\nBien cordialement,\nHugo"

    # Attacher les fichiers avec les chemins absolus
    mail.Attachments.Add(output_filename)
    mail.Attachments.Add(output_path)

    mail.Send()
    print("✅ E-mail envoyé avec succès !")

except Exception as e:
    print(f"❌ Erreur lors de l'envoi de l'e-mail : {e}")
